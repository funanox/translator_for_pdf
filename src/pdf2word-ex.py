import glob
import os
import re

from docx import Document
from pdf2docx import parse

input_path = os.path.join(os.path.dirname(__file__), "../anderson/")
temp_path = os.path.join(os.path.dirname(__file__), "../temp/")
output_path = os.path.join(os.path.dirname(__file__), "../anderson-word-EN_debug/")


def get_filename(pdf_path):
    return os.path.splitext(os.path.basename(pdf_path))[0]


def modify_doc(document):
    # ヘッダー
    def is_header(text):
        text = text.replace("\n", "")
        if re.search(r"Preface to the Third Edition$", text) is not None:
            return True
        if re.match(r"\t?\d{1,2}\.\d+\.(\s|[A-Z]|\?|-)+", text) is not None:
            return True
        return False

    # フッター（その１）
    def is_footer(text):
        text = text.replace("\n", "")
        if re.fullmatch(r"(\t)*\d{1,3}", text) is not None:
            return True
        if re.fullmatch(r"Security Engineering(\t)+\d{1,3}(\t)+Ross Anderson", text) is not None:
            return True
        return False

    def remove_footer(table):
        # フッター（その２）
        def is_footer2(text):
            text = text.replace("\n", "")
            if re.fullmatch(r"\t\d{1,3}", text) is not None:
                return True
            if re.fullmatch(r"\tRoss Anderson", text) is not None:
                return True
            if re.fullmatch(r"\tSecurity Engineering", text) is not None:
                return True
            return False

        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if is_footer2(paragraph.text):
                            paragraph.text = paragraph.text.replace(paragraph.text, "\n")

        except Exception as e:
            pass

    ligatures = {u"\ufffd": u"\u0066\u0066\u0069", u"\u21B5": u"\u0066\u0066", u"\u270f": u"\u0066\u0066\u006c"}
    for paragraph in document.paragraphs:
        # ヘッダー
        if is_header(paragraph.text):
            paragraph.text = paragraph.text.replace(paragraph.text, "\n")
        # フッター（その１）
        if is_footer(paragraph.text):
            paragraph.text = paragraph.text.replace(paragraph.text, "\n")
        # 各パラグラフのインライン処理
        inline = paragraph.runs
        for i in range(len(inline)):
            text = inline[i].text
            # リガチャ
            for key in ligatures.keys():
                if key in text:
                    text = text.replace(key, ligatures[key])
            inline[i].text = text
    # フッター（その２）
    for table in document.tables:
        remove_footer(table)

    return document


if __name__ == '__main__':
    if not os.path.isdir(output_path):
        os.mkdir(output_path)

    # PDFをWordに変換
    pdf_paths = glob.glob(input_path + "*.pdf")
    for pdf_path in pdf_paths:
        filename = get_filename(pdf_path)
        if filename == "SEv3-ch5-7sep":
            continue
        docx_temp_path = temp_path + filename + "_before_" + ".docx"
        if os.path.isfile(docx_temp_path):
            print('{} : Already PDF2WORD !!'.format(filename))
            continue
        # convert pdf to docx
        parse(pdf_path, docx_temp_path)

    # Wordファイルの整形
    word_paths = glob.glob(temp_path + "*.docx")
    for word_path in word_paths:
        filename = get_filename(word_path).replace("_before_", "")
        print(filename)
        document = Document(word_path)
        document = modify_doc(document)
        docx_path = output_path + filename + ".docx"
        document.save(docx_path)
